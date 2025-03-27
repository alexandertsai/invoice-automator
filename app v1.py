import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

# Placeholder replacement functions
def set_run_font(run, font_name):
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

def replace_in_paragraph(paragraph, data):
    full_text = "".join(run.text for run in paragraph.runs)
    replaced_text = full_text
    for placeholder, value in data.items():
        replaced_text = replaced_text.replace(placeholder, value)

    if replaced_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        new_run = paragraph.add_run(replaced_text)
        new_run.font.name = "Arial"
        new_run.font.size = Pt(8)
        set_run_font(new_run, "Arial")

def replace_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)

# Updated generate_invoices function
def generate_invoices(template_path, data_file, output_dir):
    try:
        df = pd.read_csv(data_file)    # Load the CSV file into a DataFrame

        for _, row in df.iterrows():
            # Dynamically create the data dictionary
            data = {f"{{{col}}}": str(val) for col, val in row.items()}

            new_doc = Document(template_path)  # Create a new document from the template
            replace_placeholders(new_doc, data)  # Replace placeholders with data

            output_path = os.path.join(output_dir, f"{data['{invoice_no}']}.docx")
            new_doc.save(output_path)  # Save the new document

        messagebox.showinfo("Success", "Invoices generated successfully!")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# GUI App
def main():
    def select_template():
        filepath = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if filepath:
            template_path.set(filepath)

    def select_data_file():
        filepath = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if filepath:
            data_file_path.set(filepath)

    def select_output_dir():
        directory = filedialog.askdirectory()
        if directory:
            output_dir_path.set(directory)

    def generate():
        if not template_path.get() or not data_file_path.get() or not output_dir_path.get():
            messagebox.showerror("Error", "Please select all required files and directories.")
            return
        generate_invoices(template_path.get(), data_file_path.get(), output_dir_path.get())

    root = tk.Tk()
    root.title("Invoice Generator")

    # Variables
    template_path = tk.StringVar()
    data_file_path = tk.StringVar()
    output_dir_path = tk.StringVar()

    # UI Elements
    tk.Label(root, text="Word Template:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
    tk.Entry(root, textvariable=template_path, width=40).grid(row=0, column=1, padx=10, pady=5)
    tk.Button(root, text="Browse", command=select_template).grid(row=0, column=2, padx=10, pady=5)

    tk.Label(root, text="CSV Data File:").grid(row=1, column=0, padx=10, pady=5, sticky="e")
    tk.Entry(root, textvariable=data_file_path, width=40).grid(row=1, column=1, padx=10, pady=5)
    tk.Button(root, text="Browse", command=select_data_file).grid(row=1, column=2, padx=10, pady=5)

    tk.Label(root, text="Output Directory:").grid(row=2, column=0, padx=10, pady=5, sticky="e")
    tk.Entry(root, textvariable=output_dir_path, width=40).grid(row=2, column=1, padx=10, pady=5)
    tk.Button(root, text="Browse", command=select_output_dir).grid(row=2, column=2, padx=10, pady=5)

    tk.Button(root, text="Generate Invoices", command=generate).grid(row=3, column=1, pady=20)

    root.mainloop()

if __name__ == "__main__":
    main()

