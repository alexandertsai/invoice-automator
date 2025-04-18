# Your CC and Rental csv files should be in your Downloads folder. Do not adjust their names after downloading from excel.

import os
import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime

# Placeholder replacement functions
def set_run_font(run, font_name):
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

def replace_in_paragraph(paragraph, data):
    full_text = "".join(run.text for run in paragraph.runs)
    replaced_text = full_text
    for placeholder, value in data.items():
        replaced_text = replaced_text.replace(placeholder, str(value))

    if replaced_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        new_run = paragraph.add_run(replaced_text)
        new_run.font.name = "Arial"
        new_run.font.size = Pt(8)
        set_run_font(new_run, "Arial")

def replace_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)

# Function to get next month abbreviation
def get_next_month_abbr():
    current_date = datetime.now()
    # Calculate the first day of the next month
    if current_date.month == 12:
        next_month = datetime(current_date.year + 1, 1, 1)
    else:
        next_month = datetime(current_date.year, current_date.month + 1, 1)
    
    # Get the three-letter abbreviation
    return next_month.strftime('%b')
# Function to get current month abbreviation
def get_current_month_abbr():
    return datetime.now().strftime('%b')

# Updated generate_invoices function
def generate_invoices(template_path, data_file, output_dir):
    try:
        # Ensure the output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        df = pd.read_csv(data_file)    # Load the CSV file into a DataFrame

        for _, row in df.iterrows():
            # Dynamically create the data dictionary
            data = {f"{{{col}}}": str(val) for col, val in row.items()}

            new_doc = Document(template_path)  # Create a new document from the template
            replace_placeholders(new_doc, data)  # Replace placeholders with data

            output_path = os.path.join(output_dir, f"{data['{invoice_no}']}.docx")
            new_doc.save(output_path)  # Save the new document

        messagebox.showinfo("Success", "Invoices generated successfully!")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# Functions for specific invoice types
def generate_rent_invoices():
    template_path = "/path/to/your/rent-template.docx"  # Update with your actual template path
    data_file = "/path/to/your/rent-data.csv"  # Update with your actual data file path
    
    # Create output directory with next month
    month_abbr = get_next_month_abbr()
    output_dir = f"/path/to/{month_abbr}Rental"
    
    generate_invoices(template_path, data_file, output_dir)

def generate_utilities_invoices():
    template_path = "/path/to/your/utilities-template.docx"  # Update with your actual template path
    data_file = "/path/to/your/utilities.csv" # Update with your actual data file path
    
    # Create output directory with current month
    month_abbr = get_current_month_abbr()
    output_dir = f"/path/to/{month_abbr}CC"
    
    generate_invoices(template_path, data_file, output_dir)

# GUI App
def main():
    root = tk.Tk()
    root.title("Invoice Generator")
    root.geometry("300x150")  # Set window size

    # Create a frame for buttons
    frame = tk.Frame(root)
    frame.pack(expand=True)

    # Create the two buttons
    rent_button = tk.Button(frame, text="Generate Rent Invoices", 
                           command=generate_rent_invoices,
                           width=20, height=2)
    rent_button.pack(pady=10)

    utilities_button = tk.Button(frame, text="Generate Utilities Invoices", 
                                command=generate_utilities_invoices,
                                width=20, height=2)
    utilities_button.pack(pady=10)

    root.mainloop()

if __name__ == "__main__":
    main()